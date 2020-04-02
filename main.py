import os
import csv
import docx
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from fpdf import FPDF, HTMLMixin


class GradeSheet():
    """
    >>> gs = GradeSheet()
    >>> gs.getFileCSV('E://volon/files/v.csv')
    'file exists'
    >>> gs._pathFileCSV
    'E://volon/files/v.csv'
    >>> gs.getDict()
    {...}
    """

    _pathFileCSV = ''
    _dirForSave = ''

    def getFileCSV(self, path):
        if os.path.exists(path):
            self._pathFileCSV = path
            return 'file exists'
        return 'file not exists'

    def getDirForSave(self, path):
        if os.path.isdir(path):
            self._dirForSave = path
            return 'path is not dir'
        return 'path is dir'

    def _getDict(self):
        if self._pathFileCSV == '':
            return 'path is empty'
        dict_person = {}

        with open(self._pathFileCSV, newline='') as csvfile:
            reader = csv.reader(csvfile, delimiter=';')
            person_inform = [[i[11], f'{i[8]} {i[9]} {i[10]}', i[16], i[17]] for i in reader]
            # print(person_inform)
            for row in person_inform:
                if row[3] == '1':
                    if dict_person.get(row[0]):
                        if dict_person.get(row[0]).get(row[1]):
                            dict_person[row[0]][row[1]][0] += 1
                        else:
                            dict_person[row[0]][row[1]] = [1, row[2]]
                    else:
                        dict_person[row[0]] = {row[1]: [1, row[2]]}

        return dict_person

    def createDocs(self, dict_person, stamp, date_doc):

        for faculty in dict_person:

            # формируем список с именами волонтеров
            sort_person_list = [i for i in dict_person[faculty]]
            # сортируем его
            sort_person_list.sort()

            # на основе отсортированного списка формируем новый словарь, с упорядоченным значением имен
            sort_dict = {}
            for name in sort_person_list:
                sort_dict[name] = dict_person[faculty][name]

            # Создаем документ
            doc = docx.Document()

            # Настраиваем отступы
            section = doc.sections[0]
            # section.page_height = Mm(297)
            # section.page_width = Mm(210)
            section.left_margin = Mm(20)
            section.right_margin = Mm(15)
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)

            # Заголовок
            doc.add_heading('ОЦЕНОЧНАЯ ВЕДОМОСТЬ ПО ПРОЕКТУ', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # добавляем параграфы
            doc.add_paragraph('Волонтеры: олимпиадный марафон (название проекта)').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph('Сервисный проект (тип проекта)').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph('Январь – май (срок выполнения проекта)').alignment = WD_ALIGN_PARAGRAPH.CENTER

            # добавляем пурвую таблицу 2х2
            table = doc.add_table(rows=2, cols=2)
            # выравниваем ее по левому краю
            table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # настраиваем ширину столбцов (через ячейки, т.к. word не равняет ширину колонок через column)
            cells0 = table.columns[0].cells
            for i in cells0:
                i.width = 373224 * 5
            cells1 = table.columns[1].cells
            for i in cells1:
                i.width = 373224 * 7
            # применяем стиль для таблицы
            table.style = 'Table Grid'
            # заполняем таблицу
            one = table.cell(0, 0)
            one.text = ''
            one.paragraphs[0].add_run('Руководитель проекта: ').bold = True
            one.add_paragraph('ФИО ')
            one.add_paragraph('Должность ')
            two = table.cell(0, 1)
            two.add_paragraph().add_run('Протасевич Тамара Анатольевна').bold = True
            two.add_paragraph('Директор по профессиональной ориентации и работе с одаренными учащимися')
            table.cell(1, 0).text = 'Образовательная программа'
            table.cell(1, 1).text = faculty

            # пропускаем строчку
            doc.add_paragraph('')

            # добавляем таблицу 1x4
            tableTwo = doc.add_table(rows=1, cols=4)
            # применяем стиль для таблицы
            tableTwo.style = 'Table Grid'
            # заполняем заголовки для страницы
            tableTwo.cell(0, 0).text = 'ФИО'
            tableTwo.cell(0, 1).text = 'Курс'
            tableTwo.cell(0, 2).text = 'Оценка по 10-балльной шкале'
            tableTwo.cell(0, 3).text = 'Количество ЗЕ за проект'

            # добавляем строки таблице и заполняем содержимым переданного словарая с волонтерами
            for person in sort_dict:

                r = tableTwo.add_row()
                c = r.cells
                c[0].text = person
                c[1].text = sort_dict[person][1]
                c[2].text = '10'
                grade = sort_dict[person][0]
                if grade < 3:
                    value = '1'
                elif grade == 3:
                    value = '2'
                else:
                    value = '3'
                c[3].text = value

            # задаем ширину столбцов, где 373224 количество EMU в одном сантиметре
            cells0 = tableTwo.columns[0].cells
            for i in cells0:
                i.width = 373224 * 10
            cells1 = tableTwo.columns[1].cells
            for i in cells1:
                i.width = 373224 * 4
            # cells2 = tableTwo.columns[2].cells
            # for i in cells2:
            #     i.width = 373224 * 2
            # cells3 = tableTwo.columns[3].cells
            # for i in cells3:
            #     i.width = 373224 * 2

            doc.add_paragraph()
            doc.add_paragraph()

            tableThree = doc.add_table(rows=2, cols=2)
            # Вставляем дату
            tableThree.cell(0, 0).add_paragraph(f'Дата заполнения: {date_doc}')
            tableThree.cell(0, 1).add_paragraph('Протасевич Т.А.').alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # tableThree.cell(1, 0).text = 'Оценка по 10-балльной шкале'
            if stamp:
                tableThree.cell(1, 1).text = ''
                stamp = tableThree.cell(1, 1).paragraphs[0]
                stamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                stamp.add_run().add_picture('123.png')

            # путь сохранения файла
            path_save = faculty.replace('"', '').replace(':', '')
            a = f'{self._dirForSave}/{path_save}.docx'
            # print(a) # тестовый вывод

            # сохраняем созданный документ
            doc.save(a)

    def createPDF(self, dict_person, stamp, date_doc):

        for faculty in dict_person:

            # формируем список с именами волонтеров
            sort_person_list = [i for i in dict_person[faculty]]
            # сортируем его
            sort_person_list.sort()

            # на основе отсортированного списка формируем новый словарь, с упорядоченным значением имен
            sort_dict = {}
            for name in sort_person_list:
                sort_dict[name] = dict_person[faculty][name]

            pdf = FPDF()
            pdf.add_page()
            pdf.add_font('DejaVu', '', 'DejaVuSansCondensed.ttf', uni=True)
            pdf.add_font('DejaVu', 'B', 'DejaVuSansCondensed-Bold.ttf', uni=True)
            pdf.set_font('DejaVu', 'B', 16)
            pdf.cell(0, 10, 'ОЦЕНОЧНАЯ ВЕДОМОСТЬ ПО ПРОЕКТУ', 0, 1, 'C')
            pdf.set_font('DejaVu', '', 12)
            pdf.cell(0, 10, 'Волонтеры: олимпиадный марафон(название проекта)', 0, 1, 'C')
            pdf.cell(0, 10, 'Сервисный проект (тип проекта)', 0, 1, 'C')
            pdf.cell(0, 10, 'Январь – май (срок выполнения проекта)', 0, 1, 'C')

            # Таблица 1
            pdf.cell(30, 10, '', 0, 0, 'C')
            pdf.multi_cell(65, 7, 'Руководитель проекта:\nФИО \nДолжность\n\n', 1, 'L')
            pdf.set_y(pdf.get_y() - 28)
            pdf.set_x(pdf.get_x() + 65 + 30)
            pdf.multi_cell(90, 7, 'Протасевич Тамара Анатольевна \nДиректор по профессиональной ориентации и работе с '
                                  'одаренными учащимися', 1, 'L')
            pdf.cell(30, 10, '', 0, 0, 'C')
            pdf.multi_cell(65, 7, 'Образовательная программа\n\n\n', 1, 'L')
            # устанавливаем позицию следующей мульти-ячейки
            pdf.set_xy(pdf.get_x() + 65 + 30, pdf.get_y() - 7 * 3)
            # расчет положения последней ячейки
            f = faculty + '\n\n' if 40 <= len(faculty) < 80 else faculty + '\n\n\n' if len(faculty) < 40 else faculty
            pdf.multi_cell(90, 7, f, 1, 'L')

            pdf.ln(20)

            # Таблица 2 заголовки
            w_cell1 = 90
            w_cell2 = 40
            w_cell3 = 30
            w_cell4 = 30
            h_cell = 5 * 3
            pdf.set_font('DejaVu', 'B', 12)
            pdf.multi_cell(w_cell1, 5, '\nФИО\n\n', 1, 'C')
            pdf.set_xy(pdf.get_x() + w_cell1, pdf.get_y() - h_cell)
            pdf.multi_cell(w_cell2, 5, '\nКурс\n\n', 1, 'C')
            pdf.set_xy(pdf.get_x() + w_cell1 + w_cell2, pdf.get_y() - h_cell)
            pdf.multi_cell(w_cell3, 5, 'Оценка по 10-балльной шкале', 1, 'C')
            pdf.set_xy(pdf.get_x() + w_cell1 + w_cell2 + w_cell3, pdf.get_y() - h_cell)
            pdf.multi_cell(w_cell4, 5, 'Количество ЗЕ за проект', 1, 'C')

            # Таблица 2 тело
            pdf.set_font('DejaVu', '', 12)

            # добавляем строки таблице и заполняем содержимым переданного словарая с волонтерами
            for person in sort_dict:

                pdf.cell(90, 5, person, 1, 0, 'L')
                pdf.cell(40, 5, sort_dict[person][1], 1, 0, 'C')
                pdf.cell(30, 5, '10', 1, 0, 'C')
                grade = sort_dict[person][0]
                if grade < 3:
                    value = '1'
                elif grade == 3:
                    value = '2'
                else:
                    value = '3'
                pdf.cell(30, 5, value, 1, 1, 'C')

            pdf.ln(5)

            # Вставляем дату
            pdf.cell(100, 5, f'Дата заполнения: {date_doc}', 0, 0, 'L')
            pdf.cell(90, 5, 'Протасевич Т.А.', 0, 1, 'R')

            if stamp:
                pdf.image('123.png', x=120, w=73, h=40.5)

            # путь сохранения файла
            path_save = faculty.replace('"', '').replace(':', '')
            a = f'{self._dirForSave}/{path_save}.pdf'
            # # сохраняем созданный документ
            pdf.output(a)


if __name__ == '__main__':
    import doctest

    # doctest.testmod(optionflags=+doctest.ELLIPSIS)
    gs = GradeSheet()
    gs._dirForSave = 'pdf'
    gs.getFileCSV('E://volon/files/v_origin.csv')
    # print(gs._getDict())
    # gs.createDocs(gs._getDict())
    gs.createPDF(gs._getDict(), True, '20.12.2020')
